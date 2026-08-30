#!/usr/bin/env python3
"""Build a .docx from plain or lightly structured text.

Supported light markup: ``#``/``##``/``###`` headings, blank-line separated
paragraphs, ``- ``/``* `` bullets, ``1. `` numbered items, ``| a | b |`` simple
tables (a dashed row right under the first row marks it as a header), ``---``
horizontal rule, plus inline ``**bold**`` and ``*italic*``. Russian and English
text both render correctly with the default Word fonts.
"""

from __future__ import annotations

import argparse
import re
import sys

import docx_common as common

_ORDERED_RE = re.compile(r"^(\d+)[.)]\s+(.*)$")
_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*))", re.DOTALL)


def split_row(line: str) -> list[str]:
    body = line.strip()
    if body.startswith("|"):
        body = body[1:]
    if body.endswith("|"):
        body = body[:-1]
    return [cell.strip() for cell in body.split("|")]


def is_separator_row(line: str) -> bool:
    cells = split_row(line)
    if not cells:
        return False
    return all(cell and set(cell) <= {"-", ":", " "} for cell in cells)


def parse_blocks(raw: str) -> list[dict]:
    """Translate light markup into an ordered list of block descriptors."""
    lines = raw.replace("\r\n", "\n").split("\n")
    blocks: list[dict] = []
    paragraph_lines: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = " ".join(part.strip() for part in paragraph_lines if part.strip())
        paragraph_lines.clear()
        if text:
            blocks.append({"kind": "paragraph", "text": text})

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            rows: list[list[str]] = []
            header = False
            while index < len(lines) and lines[index].strip().startswith("|"):
                candidate = lines[index].strip()
                if len(rows) == 1 and is_separator_row(candidate):
                    header = True
                else:
                    rows.append(split_row(candidate))
                index += 1
            if rows:
                blocks.append({"kind": "table", "rows": rows, "header": header})
            continue
        if set(stripped) <= {"-", "_", "*"} and len(stripped) >= 3:
            flush_paragraph()
            blocks.append({"kind": "rule"})
            index += 1
            continue
        if stripped.startswith("#"):
            flush_paragraph()
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append({
                "kind": "heading",
                "level": max(1, min(4, level)),
                "text": stripped[level:].strip(),
            })
            index += 1
            continue
        if stripped[:2] in ("- ", "* "):
            flush_paragraph()
            blocks.append({"kind": "list_item", "list": "bullet",
                           "text": stripped[2:].strip()})
            index += 1
            continue
        ordered = _ORDERED_RE.match(stripped)
        if ordered:
            flush_paragraph()
            blocks.append({"kind": "list_item", "list": "numbered",
                           "text": ordered.group(2).strip()})
            index += 1
            continue
        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    return blocks


def add_rich_text(paragraph, text: str) -> None:
    """Write *text* into *paragraph*, honouring **bold** and *italic* markers."""
    for token in _TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)


def add_rule(document) -> None:
    """Append a thin horizontal rule as a bottom-bordered empty paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def add_table(document, block: dict) -> None:
    rows = block["rows"]
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    try:
        table.style = "Table Grid"
    except Exception:  # noqa: BLE001 — an unavailable style must not break the export
        pass
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            value = row[col_index] if col_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            add_rich_text(paragraph, value)
            if block["header"] and row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def render(document, blocks: list[dict], title: str) -> None:
    if title.strip():
        document.add_heading(title.strip(), 0)
    for block in blocks:
        kind = block["kind"]
        if kind == "heading":
            document.add_heading(block["text"], block["level"])
        elif kind == "paragraph":
            add_rich_text(document.add_paragraph(), block["text"])
        elif kind == "list_item":
            style = "List Number" if block["list"] == "numbered" else "List Bullet"
            try:
                paragraph = document.add_paragraph(style=style)
            except KeyError:
                paragraph = document.add_paragraph()
            add_rich_text(paragraph, block["text"])
        elif kind == "table":
            add_table(document, block)
        elif kind == "rule":
            add_rule(document)


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Build a .docx from text.")
    parser.add_argument("--out", required=True,
                        help="output .docx file name, created under the skill "
                             "state dir outputs/")
    parser.add_argument("--in", dest="src", default="", help="text/markdown source file")
    parser.add_argument("--text", default="", help="inline text (alternative to --in)")
    parser.add_argument("--title", default="",
                        help="document title rendered at the top and in the properties")
    parser.add_argument("--author", default="", help="optional document author")
    args = parser.parse_args(argv)

    if bool(args.src.strip()) == bool(args.text.strip()):
        common.fail("provide exactly one of --in <file> or --text <string>")

    if args.src.strip():
        source = common.resolve_path(args.src, must_exist=True)
        try:
            raw = source.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            common.fail(f"cannot read source text: {exc}")
    else:
        raw = args.text

    blocks = parse_blocks(raw)
    if not blocks:
        common.fail("nothing to write: the input text is empty")

    out_path = common.resolve_output_path(args.out)
    if out_path.suffix.lower() != ".docx":
        out_path = out_path.with_suffix(".docx")

    common.require("docx", "python-docx")
    import docx

    document = docx.Document()
    try:
        render(document, blocks, args.title)
    except SystemExit:
        raise
    except Exception as exc:  # noqa: BLE001 — report layout failures honestly
        common.fail(f"failed to build the document: {exc}")

    try:
        props = document.core_properties
        if args.title.strip():
            props.title = args.title.strip()
        if args.author.strip():
            props.author = args.author.strip()
    except Exception:  # noqa: BLE001 — properties are a nicety, never a failure
        pass

    try:
        document.save(str(out_path))
    except Exception as exc:  # noqa: BLE001 — a failed save must be explicit
        common.fail(f"cannot write {out_path}: {exc}")

    counts: dict[str, int] = {}
    for block in blocks:
        counts[block["kind"]] = counts.get(block["kind"], 0) + 1
    common.emit({
        "status": "ok",
        "file": str(out_path),
        "name": out_path.name,
        "bytes": out_path.stat().st_size,
        "title": args.title,
        "blocks": counts,
    })
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
